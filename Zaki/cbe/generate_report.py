import requests
from docx import Document
from docx.shared import Inches
import logging
from PIL import Image
import io

# Set up logging
logging.basicConfig(filename="zaki.log", level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

API_URL = "https://router.huggingface.co/featherless-ai/v1/chat/completions"
HEADERS = {
    "Authorization": "Bearer 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",
}

def generate_daily_report(data: dict) -> str:
    prompt = (
        "Generate a professional Concise daily site report for a construction project. "
        "Include sections for Date, Weather, Manpower, Progress, Safety, and Photo Summary. "
        "Use a formal tone and clear structure. Format as plain text with section headers in the format '## Section Name'. "
        "Do not add any next step or optional sections and remove any \'*\' do not create any bold text at all"
        f"Input data: {data}"
    )

    payload = {
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.6,
        "top_p": 0.7,
        "model": "google/gemma-3-4b-it"
    }

    try:
        logger.info("Sending request to Hugging Face API")
        response = requests.post(API_URL, headers=HEADERS, json=payload)
        response.raise_for_status()
        result = response.json()
        if not result.get("choices") or not result["choices"][0].get("message") or not result["choices"][0]["message"].get("content"):
            raise ValueError("Invalid API response: missing choices or content")
        logger.info("Received valid response from Hugging Face API")
        return result["choices"][0]["message"]["content"]
    except requests.RequestException as e:
        logger.error(f"Hugging Face API request failed: {str(e)}")
        raise Exception(f"Failed to generate report: {str(e)}")
    except ValueError as e:
        logger.error(f"Invalid API response: {str(e)}")
        raise Exception(f"Failed to parse API response: {str(e)}")

def generate_daily_report_doc(report_text: str, output_path: str, date: str, photo_path: str = None):
    try:
        logger.info(f"Generating Word document: {output_path}")
        doc = Document()
        doc.add_heading(f"Daily Site Report - {date}", 0)
        
        # Split report into sections and format
        sections = report_text.split("\n\n")
        photo_summary_found = False
        for section in sections:
            if section.strip():
                if section.startswith("##"):
                    section_title = section.strip("## ").strip()
                    doc.add_heading(section_title, level=2)
                    if section_title.lower() == "photo summary":
                        photo_summary_found = True
                else:
                    doc.add_paragraph(section.strip())
        
        # Always append photo at the end if provided
        if photo_path:
            try:
                # Validate image
                with Image.open(photo_path) as img:
                    img.verify()  # Check if image is valid
                # Reopen image for adding to document
                doc.add_heading("Photo", level=2)
                doc.add_picture(photo_path, width=Inches(4.0))
                logger.info(f"Added photo to document: {photo_path}")
            except Exception as e:
                logger.error(f"Failed to add photo {photo_path} to document: {str(e)}")
                doc.add_paragraph(f"Error: Could not embed photo ({str(e)})")
        
        doc.save(output_path)
        # in-plan: adding the output file to a postgreSQL database.
        logger.info(f"Word document saved successfully: {output_path}")
    except Exception as e:
        logger.error(f"Failed to generate Word document: {str(e)}")
        raise Exception(f"Failed to generate Word document: {str(e)}")
